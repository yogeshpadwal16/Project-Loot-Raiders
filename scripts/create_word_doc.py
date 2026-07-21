import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()
    
    # Page Setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(11)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("🚀 Loot Raiders - Cloud VPS & PWA Deployment Blueprint")
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph("This document contains all details, credentials, management commands, and installation instructions for the Project Loot Raiders cloud deployment and Mobile PWA (Progressive Web App).").paragraph_format.space_after = Pt(24)
    
    # Section 1: Server Info
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. Cloud Server Details")
    h1_run.font.size = Pt(14)
    h1_run.font.bold = True
    h1_run.font.color.rgb = RGBColor(51, 51, 51)
    
    table1 = doc.add_table(rows=5, cols=2)
    table1.style = 'Light Shading Accent 1'
    
    data1 = [
        ("Public IP Address", "92.4.70.19"),
        ("SSH Username", "ubuntu"),
        ("Region / AD", "ap-mumbai-1 (India West - Mumbai) / AD-1"),
        ("Private Key File Path", "C:\\Users\\yoges\\Downloads\\ssh-key-2026-07-18.key"),
        ("Hardware Profile", "VM.Standard.E2.1.Micro (Always Free, 1 OCPU, 1 GB RAM + 4 GB Swap Space)")
    ]
    
    for i, (label, val) in enumerate(data1):
        table1.rows[i].cells[0].text = label
        table1.rows[i].cells[1].text = val
        table1.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Section 2: Live URLs
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. Web & API Links")
    h2_run.font.size = Pt(14)
    h2_run.font.bold = True
    h2_run.font.color.rgb = RGBColor(51, 51, 51)
    
    table2 = doc.add_table(rows=4, cols=2)
    table2.style = 'Light Shading Accent 1'
    
    data2 = [
        ("Control Panel Dashboard", "http://92.4.70.19:5555/"),
        ("API Status Check", "http://92.4.70.19:5555/api/status"),
        ("Deals Stream Feed", "http://92.4.70.19:5555/api/deals"),
        ("Telegram Channel", "https://t.me/LootRaidersDeals")
    ]
    
    for i, (label, val) in enumerate(data2):
        table2.rows[i].cells[0].text = label
        table2.rows[i].cells[1].text = val
        table2.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Section 3: Commands
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("3. SSH Management Commands")
    h3_run.font.size = Pt(14)
    h3_run.font.bold = True
    h3_run.font.color.rgb = RGBColor(51, 51, 51)
    
    p = doc.add_paragraph("Open PowerShell or Command Prompt on your local computer and run these commands to manage your scraper daemon:")
    p.paragraph_format.space_after = Pt(6)
    
    commands = [
        ("Log in to the server", 'ssh -i "C:\\Users\\yoges\\Downloads\\ssh-key-2026-07-18.key" ubuntu@92.4.70.19'),
        ("Check scraper process status", "pm2 status"),
        ("View real-time engine output logs", "pm2 logs loot-raiders"),
        ("Restart the scraper", "pm2 restart loot-raiders"),
        ("Stop the scraper", "pm2 stop loot-raiders")
    ]
    
    for desc, cmd in commands:
        p_desc = doc.add_paragraph()
        p_desc.add_run(f"• {desc}:").font.bold = True
        p_desc.paragraph_format.space_after = Pt(2)
        
        p_cmd = doc.add_paragraph()
        p_cmd.paragraph_format.left_indent = Inches(0.25)
        p_cmd_run = p_cmd.add_run(cmd)
        p_cmd_run.font.name = 'Consolas'
        p_cmd_run.font.size = Pt(9.5)
        p_cmd_run.font.color.rgb = RGBColor(102, 0, 102)
        p_cmd.paragraph_format.space_after = Pt(8)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Section 4: System Architecture
    h4 = doc.add_paragraph()
    h4_run = h4.add_run("4. System & Security Configurations")
    h4_run.font.size = Pt(14)
    h4_run.font.bold = True
    h4_run.font.color.rgb = RGBColor(51, 51, 51)
    
    points = [
        "Swap memory configured: A 4 GB swap file (/swapfile) has been permanently mapped to support Playwright/Chromium without OOM crashes.",
        "Operating System Firewall: Port 5555 has been inserted at rule position 5 in iptables (prior to the REJECT rule) and saved persistently.",
        "OCI Network Security: An Ingress Rule for Destination Port 5555 and Source CIDR 0.0.0.0/0 (TCP) has been successfully added to default subnet security list.",
        "Daemonization: The scraper and local web server run as a background service managed by PM2, configured to auto-start if the VPS reboots."
    ]
    
    for pt in points:
        doc.add_paragraph(f"✔ {pt}").paragraph_format.space_after = Pt(4)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Section 5: PWA Mobile Features & Installation
    h5 = doc.add_paragraph()
    h5_run = h5.add_run("5. Mobile App (PWA) Features & Installation")
    h5_run.font.size = Pt(14)
    h5_run.font.bold = True
    h5_run.font.color.rgb = RGBColor(51, 51, 51)
    
    p_pwa = doc.add_paragraph("The Control Panel has been configured as a Progressive Web App (PWA) with offline shell support, custom app icons (192px and 512px), splash screen styling, and standalone display. You can install it on any smartphone:")
    p_pwa.paragraph_format.space_after = Pt(8)
    
    # Android instructions
    p_and = doc.add_paragraph()
    p_and.add_run("📲 Android (Google Chrome)").font.bold = True
    p_and.paragraph_format.space_after = Pt(2)
    
    p_and_steps = doc.add_paragraph(
        "1. Open Chrome on your phone and visit http://92.4.70.19:5555/\n"
        "2. Wait for the 'Add Loot Raiders to Home screen' popup at the bottom, or\n"
        "3. Tap the three dots (menu) in the top-right corner and select 'Install app' or 'Add to Home screen'."
    )
    p_and_steps.paragraph_format.left_indent = Inches(0.25)
    p_and_steps.paragraph_format.space_after = Pt(8)
    
    # iOS instructions
    p_ios = doc.add_paragraph()
    p_ios.add_run("🍎 iOS / iPhone (Safari)").font.bold = True
    p_ios.paragraph_format.space_after = Pt(2)
    
    p_ios_steps = doc.add_paragraph(
        "1. Open Safari on your iPhone and visit http://92.4.70.19:5555/\n"
        "2. Tap the 'Share' icon (square with an up arrow) at the bottom.\n"
        "3. Scroll down the sharing menu options and tap 'Add to Home Screen'."
    )
    p_ios_steps.paragraph_format.left_indent = Inches(0.25)
    p_ios_steps.paragraph_format.space_after = Pt(8)
    
    # Save the file to the Desktop
    doc.save("C:\\Users\\yoges\\Desktop\\Loot_Raiders_VPS_Details.docx")
    print("Document updated successfully on the Desktop!")

if __name__ == "__main__":
    create_document()
