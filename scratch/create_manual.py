import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_heading_styled(doc, text, level, color_rgb):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    run = h.runs[0]
    run.font.name = 'Segoe UI'
    run.font.color.rgb = color_rgb
    return h

def main():
    desktop_path = "C:\\Users\\yoges\\Desktop"
    manual_path = os.path.join(desktop_path, "Automated_Deal_Matrix_Engine_User_Manual.docx")
    
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Color Palette (Dark Slate Blue and Gold Highlights)
    COLOR_PRIMARY = RGBColor(18, 24, 36)
    COLOR_SECONDARY = RGBColor(40, 116, 240)
    COLOR_DARK = RGBColor(28, 36, 56)
    
    # ==========================================
    # COVER PAGE
    # ==========================================
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(100)
    title_p.paragraph_format.space_after = Pt(12)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("LOOT RAIDERS\nAUTOMATED DEAL MATRIX ENGINE")
    run_title.font.name = 'Segoe UI'
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY
    
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(40)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub_p.add_run("A High-Performance E-Commerce Intelligence, Pricing Analysis, and Asynchronous Telegram Broadcast Platform")
    run_sub.font.name = 'Segoe UI'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = COLOR_SECONDARY
    
    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_before = Pt(150)
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = meta_p.add_run("USER & OPERATOR MANUAL\n\nAuthor: Yogesh Padwal\nVersion: 2.1 • Production Ready\nDate: July 2026")
    run_meta.font.name = 'Segoe UI'
    run_meta.font.size = Pt(11)
    run_meta.font.bold = True
    run_meta.font.color.rgb = RGBColor(120, 120, 120)
    
    doc.add_page_break()
    
    # ==========================================
    # 1. PLATFORM OVERVIEW
    # ==========================================
    add_heading_styled(doc, "1. Executive Introduction", 1, COLOR_PRIMARY)
    p = doc.add_paragraph(
        "Loot Raiders is an enterprise-grade Deal Intelligence Platform designed specifically to scan, "
        "verify, score, and broadcast online shopping deals from major Indian e-commerce sites. The platform "
        "replaces manual deal hunting with automated background scrapers that execute parallel crawler threads "
        "and process pricing anomalies in real-time."
    )
    p.runs[0].font.name = 'Segoe UI'
    p.runs[0].font.size = Pt(11)
    
    p2 = doc.add_paragraph("Key Architectural Advantages:")
    p2.runs[0].font.name = 'Segoe UI'
    p2.runs[0].font.bold = True
    
    points = [
        "Relational ACID SQLite Database schema eliminating data loss and concurrent file locks.",
        "Modular Plugin architecture supporting customized DOM scrapers for Amazon, Flipkart, Myntra, Ajio, Meesho, Tata CLiQ, and JioMart.",
        "Mathematical Deal Scoring Engine evaluating discount depths, rupee savings, historical low checks, and live telemetry popularity feedback.",
        "Anti-Fake Deal Shields blocking constant flat-priced accessory spam using peak-drop algorithms and price gates.",
        "Asynchronous Background Queue Dispatcher executing Telegram and Discord broadcasts with rate-limit backoff retries.",
        "Telegram Price Alert Bot letting channel subscribers register targets and cross-referencing membership before activation."
    ]
    for pt in points:
        li = doc.add_paragraph(style='List Bullet')
        run = li.add_run(pt)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(11)
        
    doc.add_page_break()
    
    # ==========================================
    # 2. RUNNING THE SCRAPER
    # ==========================================
    add_heading_styled(doc, "2. System Startup and Operation", 1, COLOR_PRIMARY)
    p = doc.add_paragraph("To boot up the complete matrix engine, execute the master file in your python environment:")
    p.runs[0].font.name = 'Segoe UI'
    
    code = doc.add_paragraph()
    code.paragraph_format.left_indent = Inches(0.5)
    r_code = code.add_run("python loot_scraper.py")
    r_code.font.name = 'Courier New'
    r_code.font.size = Pt(11)
    r_code.font.bold = True
    
    p = doc.add_paragraph("Command Line Options:")
    p.runs[0].font.name = 'Segoe UI'
    p.runs[0].font.bold = True
    
    li = doc.add_paragraph(style='List Bullet')
    run = li.add_run("--single-run : Executes a single scanning cycle across all configured retailers, exports static JSON feeds, and exits cleanly (perfect for GHA Actions or cron tasks).")
    run.font.name = 'Segoe UI'
    
    # ==========================================
    # 3. SETTINGS MATRIX TABLE
    # ==========================================
    add_heading_styled(doc, "3. Configuration Options (settings.json)", 1, COLOR_PRIMARY)
    p = doc.add_paragraph("The settings.json file controls the threshold parameters and credentials:")
    p.runs[0].font.name = 'Segoe UI'
    
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Light Shading Accent 1'
    
    headers = ["Key Name", "Default Value", "Description"]
    for i, h_text in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h_text
        set_cell_background(cell, "121824")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.name = 'Segoe UI'
        
    settings_data = [
        ("telegram_bot_token", "YOUR_BOT_TOKEN", "API Token obtained from @BotFather for alerts."),
        ("telegram_chat_id", "@LootRaidersDeals", "Target channel handle for broadcasts."),
        ("min_discount", "30.0", "Minimum discount percentage required to scan."),
        ("min_deal_price", "299", "Filters out cheap accessory spams (e.g. cables, covers)."),
        ("min_deal_savings", "250", "Minimum absolute rupee savings required to post."),
        ("blocklist_keywords", "['case', 'cover', ...]", "Banned title keywords to drop basic accessory items."),
        ("scoring_rules", "{...}", "Holds minimum publish score and category weights.")
    ]
    
    for row_idx, data in enumerate(settings_data, 1):
        for col_idx, text in enumerate(data):
            cell = table.cell(row_idx, col_idx)
            cell.text = text
            cell.paragraphs[0].runs[0].font.name = 'Segoe UI'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if row_idx % 2 == 0:
                set_cell_background(cell, "F2F2F2")
                
    doc.add_page_break()

    # ==========================================
    # 4. ANTI-SPAM AND RATINGS
    # ==========================================
    add_heading_styled(doc, "4. Premium Deal Filtering & Ratings", 1, COLOR_PRIMARY)
    p = doc.add_paragraph("To protect your channel's reputational growth, three defensive shields are enforced:")
    p.runs[0].font.name = 'Segoe UI'
    
    p = doc.add_paragraph("1. Peak-Drop Price Verification (Anti-Flat Prices)")
    p.runs[0].font.name = 'Segoe UI'
    p.runs[0].font.bold = True
    p = doc.add_paragraph(
        "Sellers often inflate list prices to claim fake '80% off' discounts. The Scorer prevents this by "
        "tracking the historical maximum (peak) price. A deal is only marked as a verified lowest price "
        "if the current price represents a genuine drop of at least 15% from its peak price. Constant-priced "
        "items (drop is 0%) are blocked."
    )
    p.runs[0].font.name = 'Segoe UI'
    
    p = doc.add_paragraph("2. Title Keyword Blocklist")
    p.runs[0].font.name = 'Segoe UI'
    p.runs[0].font.bold = True
    p = doc.add_paragraph(
        "Any deal matching banned keywords (e.g. 'cover', 'screen guard', 'cable') is discarded, preventing "
        "low-value accessories from cluttering the feed even if their price exceeds ₹299."
    )
    p.runs[0].font.name = 'Segoe UI'
    
    p = doc.add_paragraph("3. HTML Telegram Alerts Layout")
    p.runs[0].font.name = 'Segoe UI'
    p.runs[0].font.bold = True
    p = doc.add_paragraph(
        "Alerts are compiled in HTML, rendering tap-to-copy codes, strikethrough pricing, and "
        "deal ratings out of 10.0 with dynamic stars (e.g. ★★★★☆) for a premium aesthetic."
    )
    p.runs[0].font.name = 'Segoe UI'
    
    doc.add_page_break()
    
    # ==========================================
    # 5. TROUBLESHOOTING
    # ==========================================
    add_heading_styled(doc, "5. Diagnostics & Troubleshooting", 1, COLOR_PRIMARY)
    
    p = doc.add_paragraph("Issue: Dashboard shows 'Connection lost'")
    p.runs[0].font.name = 'Segoe UI'
    p.runs[0].font.bold = True
    p = doc.add_paragraph(
        "1. Check if the scraper server is running: Ensure python loot_scraper.py is executing.\n"
        "2. Avoid double-clicking dashboard/index.html. Always access it via http://127.0.0.1:5555/ in your browser.\n"
        "3. Check for Zombie Processes: On Windows, multiple python processes can lock port 5555. Run a cleanup task in PowerShell:\n"
        "   Stop-Process -Name python -Force"
    )
    p.runs[0].font.name = 'Segoe UI'
    
    p = doc.add_paragraph("Issue: Scraper drops Selenium Chromedriver sessions")
    p.runs[0].font.name = 'Segoe UI'
    p.runs[0].font.bold = True
    p = doc.add_paragraph(
        "1. Ensure Google Chrome is installed on the local system.\n"
        "2. Check for zombie chromedriver.exe instances blocking the system and kill them using:\n"
        "   taskkill /F /IM chromedriver.exe"
    )
    p.runs[0].font.name = 'Segoe UI'
    
    # Save document
    try:
        doc.save(manual_path)
        print(f"User manual compiled successfully on Desktop: {manual_path}")
    except Exception as e:
        print(f"Failed to compile user manual: {e}")

if __name__ == "__main__":
    main()
